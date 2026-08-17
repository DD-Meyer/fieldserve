from pathlib import Path

from docx import Document
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "CM3070_Final_Report_v3 (1).docx"
OUTPUT = ROOT / "CM3070_Final_Report_v3_reviewed.docx"
HEATMAP_OUTPUTS = ROOT / "ml_service" / "models" / "heatmap" / "outputs"


def find_paragraph(document: Document, needle: str):
    for paragraph in document.paragraphs:
        if needle in paragraph.text:
            return paragraph
    raise ValueError(f"Paragraph not found: {needle}")


def replace(document: Document, needle: str, text: str):
    paragraph = find_paragraph(document, needle)
    paragraph.text = text
    return paragraph


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def add_after(document: Document, anchor, text: str = "", style=None):
    paragraph = document.add_paragraph(text, style=style)
    anchor._p.addnext(paragraph._p)
    return paragraph


def add_picture_after(document: Document, anchor, path: Path, caption: str):
    image_paragraph = document.add_paragraph()
    image_paragraph.alignment = 1
    image_paragraph.add_run().add_picture(str(path), width=Inches(6.1))
    anchor._p.addnext(image_paragraph._p)
    caption_paragraph = document.add_paragraph(caption)
    caption_paragraph.alignment = 1
    image_paragraph._p.addnext(caption_paragraph._p)
    return caption_paragraph


def add_scheduler_table(document: Document, anchor):
    caption = add_after(
        document,
        anchor,
        "Table 5.5: Illustrative candidate-slot ranking using the implemented scoring constants.",
    )
    table = document.add_table(rows=1, cols=5)
    if document.tables and document.tables[0].style is not None:
        table.style = document.tables[0].style
    headers = ("Candidate", "Travel before", "Travel after", "Fragmentation penalty", "Score")
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    rows = (
        ("14:00", "5 min", "6 min", "0", "78"),
        ("11:15", "8 min", "7 min", "15", "55"),
        ("08:30", "15 min", "10 min", "0", "50"),
        ("16:15", "12 min", "14 min", "15", "33"),
    )
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = value
    caption._p.addnext(table._tbl)
    return table


def main():
    document = Document(SOURCE)
    document.core_properties.title = "FieldServe CRM: Final Project Report"
    document.core_properties.author = "Daryn Meyer"

    # Architecture documentation is now corrected, while rendered report figures remain outstanding.
    replace(
        document,
        "The project's ARCHITECTURE.md, reviewed directly for this revision, still documents",
        "The project's ARCHITECTURE.md was corrected on 17 August 2026 to match the implemented system. The FastAPI service is now documented as exposing churn, heat-map, and vehicle-damage vision routes; the offline artefacts include the deployed YOLO weights rather than OR-Tools parameters; the Django app diagram includes inspections; and the scheduling lifecycle now remains inside Django as deterministic slot recommendation rather than calling a removed FastAPI scheduling route.",
    )
    replace(
        document,
        "[[REPLACE: update ARCHITECTURE.md's Mermaid diagrams",
        "ARCHITECTURE.md now contains the corrected Mermaid source. [[REPLACE: render the corrected high-level, Django app-ownership, and ML data-flow diagrams to PNG or SVG and insert them as Figure 3.1.]]",
    )
    replace(
        document,
        "[[REPLACE: insert the updated Django app/data-ownership diagram",
        "The updated Django app/data-ownership diagram is now maintained in ARCHITECTURE.md and includes inspections as records owned by jobs and feeding analytics. [[REPLACE: render and insert this corrected Mermaid diagram here.]]",
    )
    replace(
        document,
        "The backend schema separates concerns across Django apps:",
        "The backend schema separates concerns across Django apps: businesses (Business, Membership, and Service, forming the multi-tenancy root), users (User/Profile and tenant-owned Customer), jobs (Job and scheduling state), analytics (derived churn and heat-map outputs), and inspections (tenant-isolated guided-walkaround image records and analysis JSON owned through Job). ARCHITECTURE.md now includes the inspections app and its jobs/analytics relationships. Tenant isolation remains request-scoped through accessible business identifiers and is exercised by the inspection test that rejects attaching an image to another business's job.",
    )

    # Deployed churn bundle evidence.
    replace(
        document,
        "[[REPLACE: confirm whether the risk_bucket_thresholds",
        "Inspection of the deployed churn_model.joblib bundle confirmed that risk_bucket_thresholds is {'high': 0.65, 'medium': 0.35}. Table 5.2 therefore uses the same thresholds as the deployed router defaults rather than a dataset-derived override.",
    )

    # Executed KDE evidence and figures.
    replace(
        document,
        "A dedicated offline notebook, 03_heatmap.ipynb, has been written and reviewed",
        "The offline 03_heatmap.ipynb notebook was executed on 17 August 2026 against 96,210 delivered Olist orders with valid Brazilian coordinates. Five-fold grid search on a fixed 5,000-point sample selected Gaussian bandwidth h=0.236 decimal degrees (approximately 26 km at the equator). The run generated the national, time-filtered, seasonal, and hotspot maps plus spatial-fold validation and hotspot tables retained under ml_service/models/heatmap/outputs. The methodology remains as specified above, but the resulting spatial-fold scores require a cautious interpretation rather than the notebook's templated claim of stability.",
    )
    kde_anchor = replace(
        document,
        "[[REPLACE: execute 03_heatmap.ipynb",
        "The executed KDE outputs are reproduced below. Rio de Janeiro (RJ) supplied the highest-ranked zip-prefix hotspot (22790; 136 bookings), and four of the five highest-ranked prefixes were also in RJ. These are observational demand concentrations rather than known synthetic generator locations.",
    )
    for filename, caption in (
        ("heatmap_bandwidth_cv.png", "Figure 4.1: Five-fold cross-validated KDE bandwidth selection."),
        ("heatmap_national.png", "Figure 4.2: National Olist demand-density map using the selected bandwidth."),
        ("heatmap_time_filter.png", "Figure 4.3: Peak-hours and off-peak demand-density comparison."),
        ("heatmap_seasonal_filter.png", "Figure 4.4: Q1 and Q3 seasonal demand-density comparison."),
        ("heatmap_hotspots_overlay.png", "Figure 4.5: Highest-volume zip-prefix hotspots overlaid on the KDE map."),
    ):
        kde_anchor = add_picture_after(document, kde_anchor, HEATMAP_OUTPUTS / filename, caption)

    replace(
        document,
        "As noted in Chapter 4.4, the KDE feature is implemented but not yet quantitatively evaluated.",
        "The KDE feature was quantitatively evaluated using the executed offline notebook described in Chapter 4.4. The selected bandwidth was 0.236 degrees. Held-out per-point log-likelihoods across latitude-based folds were -130.1303, -2.2083, -5.2447, -11.8292, and -866.4020, giving a mean of -203.1629 with standard deviation 335.0770.",
    )
    replace(
        document,
        "[[REPLACE: once the offline heat-map notebook is executed",
        "The very large fold-to-fold spread shows that the KDE does not generalise consistently to geographically separated latitude bands; this is evidence of concentrated regional structure and boundary extrapolation, not stable national spatial prediction. The proposal-stage criterion that the top three KDE peaks lie within 500 m of known generator hotspots was not testable because Olist is observational data and no ground-truth generator hotspots were defined. RQ3 is therefore answered cautiously: the maps expose actionable descriptive concentrations and respond to time/season filters, but predictive geographic validity beyond observed regions is not established.",
    )

    # Scheduler and platform test evidence.
    replace(
        document,
        "The deterministic scheduler (Chapter 4.6) is evaluated on functional-correctness criteria",
        "The deterministic scheduler (Chapter 4.6) is evaluated on functional-correctness criteria rather than machine-learning metrics, consistent with its reclassification in Chapter 3.4. Tests were run inside the Docker backend so that the production GDAL/PostGIS dependencies were present.",
    )
    replace(
        document,
        "[[REPLACE: run the scheduler and public-booking test suites",
        "The focused scheduler and public-booking run passed all 13 tests in 9.54 seconds. Seven scheduler tests covered working-hour boundaries, close-time overflow, travel-buffer conflicts, feasible gaps, distance-dominant travel time, self-exclusion during edits, and cancelled-booking handling. Six public-booking tests covered business/service discovery, customer and job creation, case-insensitive customer reuse, contact validation, and disabling public booking. The full Django suite passed all 32 tests in 16.36 seconds.",
    )
    scheduler_anchor = replace(
        document,
        "[[REPLACE: add a small worked-example scenario table",
        "Table 5.5 illustrates the implemented score, max(0, 100 - 2 x total travel minutes - fragmentation penalty), for four synthetic feasible gaps. It is an explanatory worked example, not a measured travel-time experiment. The 14:00 candidate ranks first because its 11 travel minutes and absence of a small-gap penalty produce the highest score.",
    )
    add_scheduler_table(document, scheduler_anchor)

    replace(
        document,
        "[[REPLACE: insert backend, frontend, and FastAPI automated test counts",
        "Automated platform evidence comprised 32 passing Django tests (16.36 seconds) and four successful FastAPI deployment smoke checks: /health, /version, /vision/status, and /openapi.json all returned HTTP 200. The vision status confirmed that the real deployed checkpoint yolov8n-cardd-5c2d7c9 was loaded. No frontend unit or component tests currently exist. Frontend static checks are not passing: ESLint reported 7 errors and 32 warnings, while TypeScript reported 2 errors in the Clerk sign-in screen. No API latency or load test has been measured, so no performance claim is made.",
    )

    # Correct conclusions that predated the new evidence.
    replace(
        document,
        "Consistent with the requirement not to present outstanding work as a completed result",
        "Consistent with the requirement not to present outstanding work as a completed result, the evidentiary status of each component is stated explicitly. Churn prediction and computer-vision damage detection have validation evidence; KDE mapping now has completed bandwidth selection and spatial-fold evaluation, although the high fold variance limits claims of geographic generalisation; and deterministic scheduling has 13 focused passing functional tests. Platform usability (RQ4), held-out CarDD test-set performance, frontend automated testing, and load performance remain unevaluated, with the evidence required to close those gaps listed in the appendix.",
    )

    # Keep the architecture table aligned with the corrected source.
    architecture_table = document.tables[0]
    architecture_table.cell(3, 1).text = "Django REST Framework (fieldserve_backend, :8000): users · businesses · jobs · analytics · inspections"

    # Remove appendix tasks now completed, and narrow partially completed tasks.
    replace(
        document,
        "This appendix collects every [[REPLACE: ...]] marker",
        "This appendix collects the evidence still outstanding after the verified updates made on 17 August 2026. Completed items (churn thresholds, KDE execution, scheduler/public-booking tests, and backend/FastAPI counts) have been removed rather than left as stale tasks.",
    )
    replace(
        document,
        "Chapter 3.2.1: update ARCHITECTURE.md's Mermaid diagrams",
        "Chapter 3.2.1: render the corrected ARCHITECTURE.md Mermaid diagrams and insert them as report figures (the Mermaid source itself is now current).",
    )
    for completed in (
        "Chapter 4.3: confirm whether the risk_bucket_thresholds",
        "Chapter 4.4: execute 03_heatmap.ipynb",
        "Chapter 5.3: run the scheduler and public-booking automated test suites",
        "Chapter 5.3: a worked-example scenario table",
    ):
        remove_paragraph(find_paragraph(document, completed))
    replace(
        document,
        "Chapter 5.5: backend/frontend/FastAPI automated test counts",
        "Chapter 5.5: add frontend unit/component tests and make ESLint and TypeScript pass; add API latency/load-test evidence if performance is discussed.",
    )

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
