from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "CM3070_Final_Report_v5 - reduced.docx"


def replace_paragraph(document: Document, starts_with: str, text: str) -> None:
    for paragraph in document.paragraphs:
        if paragraph.text.startswith(starts_with):
            paragraph.text = text
            return
    raise ValueError(f"Report paragraph not found: {starts_with}")


def main() -> None:
    document = Document(REPORT)
    document.core_properties.title = "FieldServe CRM: Final Project Report"
    document.core_properties.author = "Daryn Meyer"

    replace_paragraph(
        document,
        "The system retains the four-tier architecture",
        "The system retains the four-tier architecture described at proposal stage. The client tier is a React Native + Expo application with tabbed navigation, a drawer for secondary screens, guided vehicle inspection, and API-backed company and user profile settings. The application tier is a Django REST Framework project exposing a tenant-scoped REST API behind Clerk-issued JWT authentication, with Clerk Organizations representing the user's business workspace; a companion FastAPI service hosts churn, heat-map, and vehicle-damage routers. The data tier is PostgreSQL/PostGIS, with Django Business and Membership records linked to the Clerk organization ID. Onboarding creates the Clerk Organization first and then synchronizes it to Django, preventing an authenticated user from entering the application without a completed company workspace. Deployment uses Expo/EAS and Render-managed Docker services.",
    )
    replace_paragraph(
        document,
        "The backend schema separates concerns across Django apps:",
        "The backend schema separates concerns across Django apps: businesses (Business, Membership, Service, and the Clerk organization identifier, forming the multi-tenancy root), users (the Clerk-linked User/Profile and tenant-owned Customer), jobs (Job and scheduling state), analytics (derived churn and heat-map outputs), and inspections (tenant-isolated guided-walkaround image records and analysis JSON owned through Job). Tenant isolation remains request-scoped through active Membership records. Company onboarding creates the local Business only after a Clerk Organization has been created and supplies its organization ID; subsequent business and profile edits use authenticated Django PATCH endpoints.",
    )
    replace_paragraph(
        document,
        "Authentication uses Clerk, issuing JWTs validated by Django",
        "Authentication uses Clerk Organizations as the business-workspace boundary. Clerk issues the user session and organization context; Django validates the JWT, resolves the local User by Clerk subject, and restricts every business, job, customer, analytics, and inspection query to active local Membership records. During onboarding, the mobile client creates a Clerk Organization, activates it, obtains a session token, and POSTs the organization ID with the company details to Django. Django stores the organization ID on Business and creates the owner Membership. If the Clerk webhook is delayed, authenticated requests can retrieve missing profile fields from Clerk's Backend API and persist them locally. A signed-in user without an active organization and local business membership is routed back to onboarding rather than the main application.",
    )
    replace_paragraph(
        document,
        "The mobile client deploys via Expo/EAS; the Django backend",
        "The mobile client deploys via Expo/EAS; the Django backend and FastAPI ML service deploy as Docker services on Render. The ML service is configured as a private service and exposes internal health and inference routes. Because the vision dependency includes Torch and Ultralytics, the YOLO vehicle-damage and framing models are lazy-loaded only when their endpoints are first used, rather than during process startup. This reduces baseline memory usage and avoids loading both models for health checks, churn, or heat-map requests. The service is still subject to the 512 MB instance limit when both vision paths are exercised, so production monitoring and plan sizing remain deployment considerations.",
    )
    replace_paragraph(
        document,
        "This project designed, built, and partially evaluated FieldServe CRM",
        "This project designed, built, and partially evaluated FieldServe CRM, a mobile-first field service CRM for micro-SMEs integrating three machine learning components (customer churn prediction, YOLOv8-based vehicle-damage detection, and KDE-based spatial demand mapping) alongside deterministic booking-slot recommendation. Its concrete contributions are: a leakage-safe churn pipeline; a trained and evaluated YOLOv8 detector; a tenant-isolated Clerk Organization/Django Business architecture with mandatory company onboarding; editable business and user profile settings persisted through authenticated APIs; a Clerk-linked demo seeding workflow; and a transparent account of the scope change from ML-optimised scheduling to a deterministic alternative.",
    )
    replace_paragraph(
        document,
        "Platform usability (RQ4) and cross-device checks remain unevaluated.",
        "Platform usability (RQ4) and cross-device checks remain unevaluated. The implemented authentication path now requires a Clerk Organization and a matching Django Membership before the main application is available; business and profile editing paths are implemented and verified by type/lint checks, while end-user usability evidence is still outstanding. Backend system checks and the full Docker test suite passed after the organization, profile, seeding, and deployment changes. Held-out CarDD test-set evaluation, frontend unit tests for churn-risk thresholds, and API load-test measurements have been completed and are reported below. The evidence required to close the remaining gaps is listed in the appendix.",
    )
    replace_paragraph(
        document,
        "Dependence on location quality:",
        "Dependence on location quality: both the heat-map feature and the deterministic scheduler's travel-cost ranking depend on the availability and accuracy of customer coordinates, which is not guaranteed for every booking.",
    )
    replace_paragraph(
        document,
        "In-app payments and production-scale evaluation",
        "Clerk Organization invitations and membership administration, richer profile/company validation, live road-time integration, in-app payments, and production-scale evaluation are future work once the platform has real operator adoption. The current implementation establishes the organization-to-business synchronization contract but does not yet provide a full in-app organization switcher or invitation workflow.",
    )

    # Append deployment-specific evidence without disturbing the reduced report layout.
    document.add_paragraph(
        "Deployment evidence: the Render ML private service exceeded its 512 MB memory limit while the two YOLO paths were loaded at import time. The implementation was revised to lazy-load each model once, with a lock protecting concurrent first use. Docker compilation and backend validation passed; a production memory profile after redeployment remains required.",
    )
    document.add_paragraph(
        "Seed-data evidence: seed_demo now accepts --clerk-organization-id (or --organization-id), reuses a user's single Clerk-linked Business when the option is omitted, rejects ambiguous multi-business ownership, and generates collision-safe slugs for legacy local-only runs. This prevents demo customers and jobs from being attached to a second unlinked Django business.",
    )

    document.save(REPORT)


if __name__ == "__main__":
    main()