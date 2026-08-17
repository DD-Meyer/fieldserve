from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "Week 18" / "CM3070_Final_Report_v3_reviewed.docx"
OUTPUT = ROOT / "Week 18" / "CM3070_Final_Report_v4.docx"
APPLICATION_IMAGES = ROOT / "Week 18" / "Fieldserve - Application images"
ARCHITECTURE_IMAGES = ROOT / "Week 18" / "FieldServe - System Architecture"
CHURN_IMAGES = ROOT / "Week 18" / "Fieldserve - Churn results"


def find_paragraph(document: Document, exact_text: str):
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text.strip() == exact_text]
    if len(matches) != 1:
        raise ValueError(f"Expected one paragraph matching {exact_text!r}, found {len(matches)}")
    return matches[0]


def set_text(paragraph, text: str, *, highlight=None) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    if highlight is not None:
        run.font.highlight_color = highlight


def insert_after(paragraph, text: str, style: str | None = None):
    new_xml = OxmlElement("w:p")
    paragraph._p.addnext(new_xml)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._p.getparent().remove(new_paragraph._p)
    new_xml.getparent().replace(new_xml, new_paragraph._p)
    if style:
        new_paragraph.style = style
    new_paragraph.add_run(text)
    return new_paragraph


def insert_before(paragraph, text: str, style: str | None = None):
    new_paragraph = insert_after(paragraph, text, style)
    paragraph._p.addprevious(new_paragraph._p)
    return new_paragraph


def add_picture_after(document: Document, anchor, path: Path, caption: str, width: float = 6.1):
    if not path.exists():
        raise FileNotFoundError(path)
    image_paragraph = document.add_paragraph()
    image_paragraph.alignment = 1
    image_paragraph.add_run().add_picture(str(path), width=Inches(width))
    anchor._p.addnext(image_paragraph._p)
    caption_paragraph = document.add_paragraph(caption)
    caption_paragraph.alignment = 1
    image_paragraph._p.addnext(caption_paragraph._p)
    return caption_paragraph


def remove_paragraph(paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)
    paragraph._p = paragraph._element = None


def set_update_fields(document: Document) -> None:
    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def main() -> None:
    document = Document(SOURCE)

    set_text(
        find_paragraph(document, "Document version 3 — final submission"),
        "Document version 4 — literature and evidence integration update",
    )

    set_text(
        find_paragraph(
            document,
            "Sections and figures that depend on evidence not available at the time of writing are marked with a highlighted placeholder in the form [[REPLACE: ...]], stating exactly what is required and how to obtain it. These markers are never presented as results; they are collected in the closing appendix, “Outstanding Evidence Required”. Two evidentiary points are flagged explicitly here because they affect how earlier and later material in this report should be read together:",
        ),
        "Version 4 integrates the architecture, heat-map, guided-inspection, damage-report, and churn-result figures supplied in the Week 18 evidence folders. Highlighted [[REPLACE: ...]] markers now remain only where the corresponding in-body evidence is not present: the inspection-imagery policy, held-out CarDD test-set evaluation, and usability study. Unresolved frontend and non-functional checks are recorded separately in Appendix A. Two evidentiary points are flagged explicitly here because they affect how earlier and later material in this report should be read together:",
    )

    chapter_intro = find_paragraph(
        document,
        "This chapter is substantially unchanged from the preliminary submission for the strands unaffected by the scope change (2.1, 2.2, 2.5), revised for the scheduling strand (2.3) to reflect that the implemented system no longer uses machine-learning-optimised routing, and extended with a new strand (2.6) covering computer-vision damage detection, which is presented as an outstanding literature gap rather than a completed review, consistent with the no-invented-references requirement governing this revision.",
    )
    set_text(
        chapter_intro,
        "This chapter retains the preliminary submission's core strands on SME digital transformation, churn prediction, spatial demand analysis, and mobile field-work UX; revises the scheduling strand to match the implemented deterministic approach; and adds a completed review of computer-vision vehicle-damage detection. The revised synthesis connects these literatures to the implemented system while distinguishing established evidence from claims that still require project-specific evaluation.",
    )

    set_text(
        find_paragraph(document, "2.6 Computer-Vision Vehicle-Damage Detection — An Outstanding Literature Strand"),
        "2.6 Computer-Vision Vehicle-Damage Detection",
    )
    first_cv = find_paragraph(
        document,
        "The third machine learning component implemented in this project, YOLOv8-based detection and classification of visible vehicle damage (dent, scratch, crack, glass shatter, broken lamp, flat tyre) from images captured during a guided walkaround, was not part of the platform's scope at the preliminary submission stage and therefore has no corresponding literature review from that submission to revise. A rigorous review of this strand, covering automated vehicle-damage assessment, object detection architectures (particularly the YOLO family), the CarDD dataset used for training (Chapter 3.5, Chapter 4.6), and domain-transfer considerations between a general-purpose damage-detection dataset and this project's specific inspection use case, has not yet been completed and is not fabricated here.",
    )
    set_text(
        first_cv,
        "Automated vehicle-damage assessment seeks to localise and classify defects from photographs, reducing the time and inconsistency associated with manual inspection while retaining an auditable visual record. Earlier systems were constrained by small or private insurance datasets, limiting reproducibility and comparison. Wang, Li and Wu (2023) addressed this problem with CarDD, a public benchmark containing 4,000 high-resolution images and more than 9,000 annotated instances across six categories: dent, scratch, crack, glass shatter, lamp broken, and tyre flat. These classes map directly to FieldServe CRM's implemented detector, making CarDD a defensible training source, although benchmark relevance does not itself establish performance on photographs captured by this application's users.",
    )
    second_cv = find_paragraph(
        document,
        "[[REPLACE: complete a literature review strand on automated vehicle-damage detection and YOLO-family object detection, citing verified sources including at minimum the CarDD dataset paper, a YOLOv8 architecture reference, and at least one prior study applying object detection to vehicle-damage or insurance-inspection use cases; insert as Chapter 2.6 in place of this paragraph, and update the synthesis in 2.7 accordingly]]",
    )
    set_text(
        second_cv,
        "The detector follows the You Only Look Once lineage. Rather than applying a separate region-proposal stage, the original YOLO formulation predicts bounding boxes and class probabilities in a single end-to-end network, prioritising an accuracy-speed trade-off suitable for interactive use (Redmon et al., 2016). YOLOv8 extends this single-stage family with an anchor-free split head and model variants supporting detection, segmentation, and other vision tasks (Jocher, Chaurasia and Qiu, 2023). Importantly, Ultralytics has not published a formal YOLOv8 architecture paper; it is therefore treated here as versioned software, while the peer-reviewed YOLO paper supports the underlying architectural rationale. The lightweight YOLOv8n variant selected by FieldServe CRM reflects a deployment trade-off: guided feedback and damage overlays need responsive inference, but small or low-contrast defects may be harder to localise than large, visually distinct damage.",
    )
    third_cv = insert_after(
        second_cv,
        "Evidence from operational insurance settings supports the feasibility of this approach but also shows why benchmark metrics require careful interpretation. Pérez-Zarate et al. (2024) evaluated automated car-damage assessment on imagery labelled under professional appraiser supervision and reported that an ensemble of YOLOv5 detectors improved both detection performance and throughput, with comparison against YOLOv8. Their work demonstrates practical demand for scalable image-based assessment, yet its ensemble architecture and insurance-company workflow differ from FieldServe CRM's single lightweight model and operator-guided walkaround. Results cannot therefore be transferred directly between the systems.",
        "Normal",
    )
    insert_after(
        third_cv,
        "Three limitations follow from this literature. First, damage classes vary greatly in visual scale and texture, so aggregate mean average precision can conceal weak performance on dents, scratches, and cracks; per-class precision, recall, mAP@50, and mAP@50–95 are required. Second, changes in camera, lighting, distance, viewpoint, compression, and background create domain shift between CarDD and mobile field inspections. Third, a bounding box is decision support rather than a repair-cost or liability judgement. FieldServe CRM consequently presents the image, class, confidence, and location to the operator, preserves human review, and evaluates the deployed model by class in Chapter 5.2 rather than claiming autonomous assessment.",
        "Normal",
    )

    synthesis_intro = find_paragraph(
        document,
        "The four strands reviewed above (excluding the outstanding computer-vision strand, 2.6, pending completion) converge on a consistent finding: while each individual dimension of this project has a substantial body of research, their combination in a single platform serving micro-SME field service businesses has not been addressed.",
    )
    set_text(
        synthesis_intro,
        "The five technical and socio-technical strands reviewed above converge on a consistent finding: each individual dimension has an established research base, but their combination in one mobile platform for micro-SME field service businesses remains underexplored. The literature supports the component choices while also defining the limits against which FieldServe CRM must be evaluated.",
    )
    gap_two = find_paragraph(
        document,
        "Gap 2: Computer-vision-based damage assessment, integrated into a guided mobile inspection workflow for micro-SME field service operators, has not been addressed in the literature reviewed to date (pending completion of 2.6); the evaluation in Chapter 5.2 is offered as an initial empirical contribution regardless, since the model has been trained and evaluated even though the surrounding literature review is incomplete.",
    )
    set_text(
        gap_two,
        "Gap 2: Vehicle-damage research establishes that object detectors can support insurance assessment and provides public benchmarks such as CarDD, but it predominantly evaluates model performance in isolation. Less attention is given to guided capture, operator correction, evidence retention, and the integration of detections into a micro-SME service workflow. Chapter 5.2 evaluates the detector, while the implemented walkaround addresses this workflow-level gap without claiming autonomous damage valuation.",
    )
    contribution = find_paragraph(
        document,
        "This project addresses these gaps through the design, development, and evaluation of FieldServe CRM. Its contributions to date are: an empirical, methodologically transparent evaluation of churn-prediction algorithms on a repeat-purchase-dense dataset, including a documented diagnosis and correction of a degenerate-label failure mode (Chapter 5.1); a trained and evaluated YOLOv8 vehicle-damage detection model with a documented resolution/augmentation ablation (Chapter 5.2); and an architecture and implementation for KDE-based spatial demand mapping and deterministic scheduling, whose quantitative validation remains outstanding and is marked accordingly throughout this report.",
    )
    set_text(
        contribution,
        "This project addresses these gaps through the design, development, and evaluation of FieldServe CRM. Its evidenced contributions are: a methodologically transparent churn-model comparison on repeat-purchase data, including diagnosis of a degenerate-label failure mode (Chapter 5.1); a YOLOv8 vehicle-damage detector evaluated overall and by class, with a resolution and augmentation ablation (Chapter 5.2); deterministic slot recommendation supported by functional tests (Chapter 5.3); and an executed KDE demand-mapping pipeline whose unstable spatial-fold result constrains claims of geographic generalisation (Chapter 5.4). Usability, held-out vision, and non-functional evidence still required are recorded in Appendix A.",
    )

    architecture_placeholder = find_paragraph(
        document,
        "ARCHITECTURE.md now contains the corrected Mermaid source. [[REPLACE: render the corrected high-level, Django app-ownership, and ML data-flow diagrams to PNG or SVG and insert them as Figure 3.1.]]",
    )
    set_text(
        architecture_placeholder,
        "The rendered architecture evidence supplied for this revision is reproduced in Figures 3.1–3.3. Together, the diagrams show the deployed tiers, Django application ownership including inspections, and the offline-training/online-inference boundary for all three ML features.",
    )
    architecture_anchor = architecture_placeholder
    for filename, caption in (
        ("1. High-level system map.png", "Figure 3.1: High-level FieldServe system map."),
        ("3. Backend app structure.png", "Figure 3.2: Django application ownership and data relationships, including inspections."),
        ("4. ML feature data flow.png", "Figure 3.3: Offline training and online inference data flow."),
    ):
        architecture_anchor = add_picture_after(
            document, architecture_anchor, ARCHITECTURE_IMAGES / filename, caption
        )

    set_text(
        find_paragraph(
            document,
            "[[REPLACE: insert an updated database ER diagram reflecting the inspections app (tenant-scoped inspection persistence, referenced in the LATEST IMPLEMENTATION CHANGES record but not present in ARCHITECTURE.md's current Django app-structure diagram) and any other schema changes made since the preliminary submission]]",
        ),
        "Figure 3.2 provides the supplied application-level data-ownership view: inspections are owned through jobs, businesses own tenant-scoped users and jobs, and inspection outputs feed analytics. The corresponding model implementation stores each inspection against a job with creator, phase, angle, image, analysis JSON, status, and timestamps (Chapter 4.1).",
    )

    set_text(
        find_paragraph(
            document,
            "The updated Django app/data-ownership diagram is now maintained in ARCHITECTURE.md and includes inspections as records owned by jobs and feeding analytics. [[REPLACE: render and insert this corrected Mermaid diagram here.]]",
        ),
        "The updated Django app/data-ownership diagram is reproduced as Figure 3.2 and includes inspections as records owned by jobs and feeding analytics.",
    )

    heatmap_placeholder = find_paragraph(
        document,
        "[[REPLACE: capture the implemented Leaflet heat map from the running application, showing clusters and density markers, as a screenshot]]",
    )
    set_text(
        heatmap_placeholder,
        "Figure 4.6 shows the implemented mobile demand heat map with KDE density, clustered customer counts, time filters, and ranked demand zones.",
    )
    add_picture_after(
        document,
        heatmap_placeholder,
        APPLICATION_IMAGES / "Heatmap.png",
        "Figure 4.6: Implemented mobile demand heat map with KDE density and clustered customer markers.",
        3.1,
    )

    inspection_placeholder = find_paragraph(
        document,
        "[[REPLACE: insert screenshots of: the guided vehicle walkaround capture screen; the automatic vehicle-framing guidance overlay (distance/centring/clipping); a damage image with bounding-box overlays as rendered in the app; and the inspection summary breakdown screen]]",
    )
    set_text(
        inspection_placeholder,
        "Figures 4.7–4.9 provide the supplied mobile evidence for the guided capture and review workflow. Figure 4.7 combines the required-view prompt, progress indicator, vehicle outline, and automatic framing guidance. Figure 4.8 shows model detections rendered as bounding boxes with class, confidence, region, and image-area details. Figure 4.9 shows the aggregate inspection summary and before/after workflow state.",
    )
    inspection_anchor = inspection_placeholder
    for filename, caption, width in (
        ("inpsection-camera_images.jpg", "Figure 4.7: Guided walkaround capture and automatic vehicle-framing guidance.", 6.1),
        ("inspection_report-before_images.png", "Figure 4.8: Damage detections and bounding-box overlays in the mobile inspection report.", 3.1),
        ("inspection_report-before.png", "Figure 4.9: Inspection summary breakdown and completed before-service walkaround.", 3.1),
    ):
        inspection_anchor = add_picture_after(
            document, inspection_anchor, APPLICATION_IMAGES / filename, caption, width
        )

    churn_placeholder = find_paragraph(
        document,
        "[[REPLACE: export churn_curves_retailii.png (PR and ROC curves for the four model variants) and a feature-importance/coefficient figure for the deployed model from the churn notebook, and insert as Figures 5.1 and 5.2]]",
    )
    set_text(
        churn_placeholder,
        "The supplied notebook outputs are reproduced in Figures 5.1 and 5.2. The curves support the comparative metrics in Table 5.1, while the gain-based feature-importance plot shows that recent and annual booking frequency contribute most strongly to the XGBoost comparator; importance denotes predictive contribution, not causality.",
    )
    churn_anchor = add_picture_after(
        document,
        churn_placeholder,
        CHURN_IMAGES / "Screenshot 2026-08-17 134201.png",
        "Figure 5.1: Precision-recall and ROC curves for the four churn-model variants.",
    )
    add_picture_after(
        document,
        churn_anchor,
        CHURN_IMAGES / "Screenshot 2026-08-17 134216.png",
        "Figure 5.2: XGBoost gain-based feature importance on Online Retail II.",
        5.4,
    )

    figure_renumbering = (
        ("Figure 5.5: Example validation batch — deployed model predictions for the same images, for direct visual comparison with Figure 5.4.", "Figure 5.7: Example validation batch — deployed model predictions for the same images, for direct visual comparison with Figure 5.6."),
        ("Figure 5.4: Example validation batch — ground-truth labels.", "Figure 5.6: Example validation batch — ground-truth labels."),
        ("Figure 5.3: Normalised confusion matrix for the deployed model. The largest source of error for dent, scratch, and crack is under-detection against background (false negatives) rather than confusion between damage classes.", "Figure 5.5: Normalised confusion matrix for the deployed model. The largest source of error for dent, scratch, and crack is under-detection against background (false negatives) rather than confusion between damage classes."),
        ("Figure 5.2: Precision-recall curves per class for the deployed model, showing the strong separation between the three high-performing classes (glass shatter, tire flat, lamp broken) and the three weaker classes (dent, scratch, crack).", "Figure 5.4: Precision-recall curves per class for the deployed model, showing the strong separation between the three high-performing classes (glass shatter, tire flat, lamp broken) and the three weaker classes (dent, scratch, crack)."),
        ("Figure 5.1: Training and validation loss and metric curves across training for the deployed model.", "Figure 5.3: Training and validation loss and metric curves across training for the deployed model."),
    )
    for old_text, new_text in figure_renumbering:
        set_text(find_paragraph(document, old_text), new_text)
    cv_comparison = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("A clear class-level pattern is visible")
    )
    set_text(
        cv_comparison,
        cv_comparison.text.replace("Figures 5.1–5.5", "Figures 5.3–5.7").replace(
            "Figure 5.3", "Figure 5.5"
        ),
    )

    mobile_ai = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("Communicating ML outputs to non-technical users")
    )
    set_text(mobile_ai, mobile_ai.text.replace("Figure 4.3", "Figure 4.8"))

    set_text(
        find_paragraph(
            document,
            "All references are formatted in Harvard style. URLs are provided where the source is openly available online. New references required for the outstanding computer-vision literature strand (Chapter 2.6) are not included here, since inventing citation details would violate the evidentiary standard set for this revision; they should be added once that strand is completed.",
        ),
        "All references are formatted in Harvard style. URLs and DOIs are provided where the source is openly available online. The computer-vision sources added in version 4 were checked against the publisher, Crossref, arXiv, or the software maintainer's documentation on 17 August 2026.",
    )

    reference_kruber = find_paragraph(
        document,
        "Kruber, M., Lübbecke, M.E. and Parmentier, A. (2017) ‘Learning when to use a decomposition’, in Integration of AI and OR Techniques in Constraint Programming. CPAIOR 2017. Springer, pp. 202–210.",
    )
    insert_before(
        reference_kruber,
        "Jocher, G., Chaurasia, A. and Qiu, J. (2023) Ultralytics YOLOv8 (Version 8.0.0) [Computer software]. Available at: https://github.com/ultralytics/ultralytics (Accessed: 17 August 2026).",
        "Normal",
    )
    reference_perron = find_paragraph(
        document,
        "Perron, L. and Furnon, V. (2024) OR-Tools (Version 9.x). Google. Available at: https://developers.google.com/optimization (Accessed: May 2026).",
    )
    insert_before(
        reference_perron,
        "Pérez-Zarate, S.A., Corzo-García, D., Pro-Martín, J.L., Álvarez-García, J.A., Martínez-del-Amor, M.A. and Fernández-Cabrera, D. (2024) ‘Automated car damage assessment using computer vision: insurance company use case’, Applied Sciences, 14(20), 9560. doi: 10.3390/app14209560.",
        "Normal",
    )
    reference_reichheld = find_paragraph(
        document,
        "Reichheld, F.F. and Sasser, W.E. (1990) ‘Zero defections: Quality comes to services’, Harvard Business Review, 68(5), pp. 105–111.",
    )
    insert_after(
        reference_reichheld,
        "Redmon, J., Divvala, S., Girshick, R. and Farhadi, A. (2016) ‘You only look once: unified, real-time object detection’, in Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition, pp. 779–788. doi: 10.1109/CVPR.2016.91.",
        "Normal",
    )
    reference_wei = find_paragraph(
        document,
        "Wei, J.-T., Lin, S.-Y. and Wu, H.-H. (2010) ‘A review of the application of RFM model’, African Journal of Business Management, 4(19), pp. 4199–4206.",
    )
    insert_after(
        reference_wei,
        "Wang, X., Li, W. and Wu, Z. (2023) ‘CarDD: a new dataset for vision-based car damage detection’, IEEE Transactions on Intelligent Transportation Systems, 24(7), pp. 7202–7214. doi: 10.1109/TITS.2023.3258480.",
        "Normal",
    )

    appendix_heading = find_paragraph(document, "Appendix: Outstanding Evidence Required")
    set_text(appendix_heading, "Appendix A: Evidence Register and Outstanding Work")
    green_key = find_paragraph(document, "Green – Done")
    set_text(green_key, "A.1 Verified evidence as at 17 August 2026")
    green_key.style = "Heading 2"
    red_key = find_paragraph(document, "Red – Do not have and needs to be done")
    set_text(
        red_key,
        "Completed items are listed with their reproducible evidence; remaining items are highlighted and must not be reported as completed results.",
    )
    register_intro = find_paragraph(
        document,
        "This appendix collects the evidence still outstanding after the verified updates made on 17 August 2026. Completed items (churn thresholds, KDE execution, scheduler/public-booking tests, and backend/FastAPI counts) have been removed rather than left as stale tasks.",
    )
    set_text(
        register_intro,
        "The evidence below is drawn from REPORT_TEST_EVIDENCE.md and the version 4 literature check. Commands, counts, limitations, and unresolved checks are retained so that claims in Chapters 2–5 remain auditable.",
    )
    completed_literature = find_paragraph(
        document,
        "Chapter 2.6: complete a literature review strand on automated vehicle-damage detection and YOLO-family object detection, with verified citations (CarDD dataset paper, YOLOv8 reference, at least one prior damage/insurance-inspection detection study).",
    )
    set_text(
        completed_literature,
        "Completed — Chapter 2.6 now reviews CarDD, the peer-reviewed YOLO lineage, YOLOv8 software, and an operational insurance-company study; four verified references were added.",
    )
    completed_literature.style = "List Paragraph"
    cursor = completed_literature
    for item in (
        "Completed — Backend test suite: 32 tests passed in Docker; scheduler and public-booking focused suite: 13 tests passed.",
        "Completed — FastAPI smoke checks: /health, /version, /vision/status, and /openapi.json returned HTTP 200; deployed vision model reported as yolov8n-cardd-5c2d7c9.",
        "Completed with limitation — KDE notebook used 96,210 located records and selected bandwidth 0.236 degrees by five-fold cross-validation; spatial-fold log-likelihood was -203.1629 ± 335.0770 per point, indicating unstable geographic generalisation.",
        "Completed with limitation — Churn risk thresholds are High ≥ 0.65 and Medium ≥ 0.35; the report retains the leakage-safe held-out model comparison and must not infer results from the empty churn_results_retailii.csv file.",
    ):
        cursor = insert_after(cursor, item, "List Paragraph")

    completed_architecture = find_paragraph(
        document,
        "Chapter 3.2.1: render the corrected ARCHITECTURE.md Mermaid diagrams and insert them as report figures (the Mermaid source itself is now current).",
    )
    set_text(
        completed_architecture,
        "Completed — Figures 3.1–3.3 reproduce the supplied high-level system, Django app-ownership, and ML data-flow diagrams.",
    )
    completed_schema = find_paragraph(
        document,
        "Chapter 3.2 / 4.1: updated database ER diagram and Django app/data-ownership diagram including the inspections app.",
    )
    set_text(
        completed_schema,
        "Completed — Figure 3.2 documents the supplied app/data-ownership relationships including inspections; Chapter 4.1 records the persisted inspection fields.",
    )
    completed_cursor = completed_schema
    for item in (
        "Completed — Figure 4.6 reproduces the supplied implemented mobile heat map with clusters and density markers.",
        "Completed — Figures 4.7–4.9 reproduce the supplied guided capture, framing guidance, damage overlays, and inspection summary evidence.",
        "Completed — Figures 5.1–5.2 reproduce the supplied churn PR/ROC curves and XGBoost feature-importance output.",
    ):
        completed_cursor = insert_after(completed_cursor, item, "List Paragraph")

    for completed_task in (
        "Chapter 4.4: screenshot of the implemented Leaflet heat map showing clusters and density markers.",
        "Chapter 4.5: screenshots of the guided walkaround capture screen, the framing-guidance overlay, a damage image with bounding-box overlays as rendered in the app, and the inspection summary breakdown screen.",
        "Chapter 5.1: churn_results_retailii.csv, provided for this revision, is empty; export churn_curves_retailii.png and a feature-importance/coefficient figure for the deployed churn model directly from the training notebook instead.",
    ):
        remove_paragraph(find_paragraph(document, completed_task))

    first_outstanding = find_paragraph(
        document,
        "Chapter 3.7: an explicit inspection-imagery consent and retention policy.",
    )
    outstanding_heading = insert_before(first_outstanding, "A.2 Outstanding evidence", "Heading 2")
    for paragraph in document.paragraphs:
        if paragraph._p is outstanding_heading._p:
            continue
        if paragraph._p.getparent() is not None and paragraph.style.name == "List Paragraph" and paragraph._p.getprevious() is not None:
            if paragraph.text.startswith("Chapter "):
                for run in paragraph.runs:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # Correct stale internal section references in the literature chapter.
    kde_implementation = next(p for p in document.paragraphs if p.text.startswith("The retail application of heat maps"))
    set_text(kde_implementation, kde_implementation.text.replace("(Chapter 4.5)", "(Chapter 4.4)"))

    properties = document.core_properties
    properties.title = "FieldServe CRM: Final Project Report"
    properties.subject = "Version 4 literature review and Week 18 evidence integration"
    properties.comments = "Chapter 2.6 completed and supplied Week 18 figures integrated on 17 August 2026."
    set_update_fields(document)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
